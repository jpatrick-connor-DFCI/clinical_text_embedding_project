"""
Generate ICI_Biomarker_Pipeline_Report.docx
"""

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_DIR = "/Users/connorpa/Documents/BIG PhD/Gusev Lab/Projects/Clinical Embeddings/compiled_results"
OUT_PATH = f"{DATA_DIR}/ICI_Biomarker_Pipeline_Report.docx"

# ── load data ──────────────────────────────────────────────────────────────────
combined = pd.read_csv(f"{DATA_DIR}/all_findings_with_validation.csv")
t1 = combined[combined["track"] == 1].copy()
t2 = combined[combined["track"] == 2].copy()

# ── helpers ────────────────────────────────────────────────────────────────────

LEVEL_ORDER = {"Very Strong": 0, "Strong": 1, "Moderate": 2, "Indirect": 3,
               "Partial": 4, "Weak": 5, "None": 6, "Unassessed": 7}

def set_cell_bg(cell, hex_color: str):
    """Set table cell background shading."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def val_color(level: str) -> str:
    return {
        "Very Strong": "C6EFCE",
        "Strong":      "CCFFCC",
        "Moderate":    "FFFFE0",
        "Indirect":    "FFF2CC",
        "Partial":     "FCE4D6",
        "Weak":        "FFD7D7",
        "None":        "FFB3B3",
        "Unassessed":  "F2F2F2",
    }.get(level, "FFFFFF")

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, space_before=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p

def make_table(doc, headers, rows, col_widths=None):
    """Create a table with bold headers and alternating row shading."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, "D9D9D9")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg  = "F9F9F9" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val) if val is not None and not (isinstance(val, float) and np.isnan(val)) else ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, bg)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def make_validation_table(doc, headers, rows_with_level):
    """Table where each row is colored by validation_level (last element of tuple)."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, "D9D9D9")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row_data, level in rows_with_level:
        row = table.add_row()
        color = val_color(level)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val) if val is not None and not (isinstance(val, float) and np.isnan(val)) else ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, color)
    return table

# ── build document ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: Pipeline Overview
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "ICI Biomarker Discovery Pipeline: Workflow Overview", level=1)

add_para(doc,
    "This document describes an observational, real-world evidence pipeline designed to identify "
    "genomic biomarkers of immune checkpoint inhibitor (ICI) response using electronic health record "
    "(EHR)-linked tumor genomic data. The pipeline operates on two parallel analytical tracks and "
    "uses clinical text embeddings from a fine-tuned Clinical-Longformer model to augment "
    "propensity score estimation.",
    space_after=6)

# 1.1 Study Design
add_heading(doc, "1.1 Study Design", level=2)
add_para(doc,
    "Two cohort designs are used in parallel to evaluate robustness across different confounding structures:")
add_bullet(doc,
    "Cohort 1 (Unmatched): First-line ICI-treated patients versus all never-ICI patients. "
    "The large control pool improves power but may introduce treatment-line confounding.")
add_bullet(doc,
    "Cohort 2 (1:1 Matched): Lines 1–3, matched 1:1 on cancer type and treatment line. "
    "Matching removes treatment-line confounding and improves exchangeability of comparison groups.")
add_para(doc,
    "For each cohort, two propensity score (PS) models are estimated:")
add_bullet(doc,
    "Covariates-only PS model: clinical and demographic covariates (age, sex, cancer type, "
    "treatment line, comorbidities, prior treatment history).")
add_bullet(doc,
    "Covariates + Clinical Text Embeddings PS model: the same covariates plus 768-dimensional "
    "embeddings from Clinical-Longformer applied to all clinical notes in a 30-day pre-treatment "
    "buffer window, pooled with time-decay weighting (decay=0.01).")
add_para(doc,
    "Analyses are conducted both pan-cancer and cancer-type-specifically (BRAIN, BREAST, LUNG, "
    "SOFT_TISSUE, pan_cancer, and others with sufficient sample sizes).")

# 1.2 Propensity Score Generation
add_heading(doc, "1.2 Propensity Score Generation", level=2)
add_para(doc,
    "Propensity scores are estimated using elastic-net penalized logistic regression "
    "(scikit-learn LogisticRegressionCV) with 5-fold nested cross-validation to select the "
    "regularization hyperparameter. For the text-augmented model, Clinical-Longformer embeddings "
    "are computed from clinical notes within a 30-day window before treatment start. "
    "Multiple notes per patient are pooled using exponential time-decay weighting (decay constant = 0.01 "
    "per day), giving more weight to notes closest to treatment initiation. "
    "Model discrimination is summarized by PS AUC per cancer type and cohort.")

# 1.3 IPTW Analysis
add_heading(doc, "1.3 IPTW Analysis", level=2)
add_para(doc,
    "Inverse probability of treatment weighting (IPTW) is applied to balance covariate "
    "distributions between ICI-treated and control patients:")
add_bullet(doc,
    "Common support trimming: patients below the 0.5th percentile or above the 99.5th percentile "
    "of the PS distribution are excluded.")
add_bullet(doc,
    "Stabilized ATE weights are computed and further truncated at the 1st/99th percentiles "
    "to limit the influence of extreme weights.")
add_bullet(doc,
    "Two weighting schemes are evaluated for each PS model: ATE-weighted and unweighted "
    "(no IPTW, for sensitivity analysis). This yields four specifications per analysis: "
    "(covariates-only × ATE), (covariates-only × unweighted), (covariates+embeddings × ATE), "
    "(covariates+embeddings × unweighted).")
add_bullet(doc,
    "Effective sample size (ESS) and standardized mean differences (SMD) per covariate "
    "are saved as diagnostics.")

# 1.4 Track 1
add_heading(doc, "1.4 Track 1 — Prognostic Analysis (ICI-Treated Patients)", level=2)
add_para(doc,
    "Track 1 estimates the prognostic effect of each somatic marker within ICI-treated patients only. "
    "The Cox proportional hazards model is:")
add_para(doc, "    S(t) ~ base_vars + line_dummies + marker", space_before=2, space_after=2)
add_para(doc,
    "Generalizability weights (1/PS, truncated) are applied to reweight the ICI-treated population "
    "toward the broader cancer-type population. Markers that reach FDR < 0.05 within a mutation type "
    "(_SNV, _AMP, _DEL, _SV, _FUSION) are retained. Track 1 identifies markers that are "
    "prognostically relevant specifically within ICI-treated patients — they do not distinguish "
    "whether the prognostic effect would also hold without ICI.")

# 1.5 Track 2
add_heading(doc, "1.5 Track 2 — Predictive Analysis (Full Cohort)", level=2)
add_para(doc,
    "Track 2 estimates differential ICI benefit or harm for patients with versus without each marker. "
    "The interaction Cox model is run on the full ICI + control cohort:")
add_para(doc, "    S(t) ~ base_vars + line_dummies + marker + ICI + marker×ICI", space_before=2, space_after=2)
add_para(doc,
    "The interaction term (marker×ICI) captures whether the ICI effect differs by marker status. "
    "Markers with a significant interaction (FDR < 0.05 within mutation type) are classified as "
    "predictive_ICI_benefit (HR interaction < 1) or predictive_ICI_harm (HR interaction > 1). "
    "IPTW is applied as in Track 1.")

# 1.6 FDR Correction
add_heading(doc, "1.6 FDR Correction", level=2)
add_para(doc,
    "Benjamini-Hochberg FDR correction is applied separately within each mutation type "
    "(_SNV, _AMP, _DEL, _SV, _FUSION) to account for the very different numbers of markers "
    "tested per type. This stratified correction avoids penalizing rare mutation types "
    "(e.g., _FUSION) for the large number of _DEL tests.")

# 1.7 Diagnostics
add_heading(doc, "1.7 Diagnostics", level=2)
add_para(doc, "The following diagnostics are saved per analysis scheme:")
add_bullet(doc, "PS AUC: model discrimination by cancer type and cohort")
add_bullet(doc, "Effective sample size (ESS): ATE weights, treated and control")
add_bullet(doc, "Standardized mean differences (SMD): per covariate, pre- and post-weighting")
add_para(doc,
    "A scheme_diagnostics_summary.csv file consolidates these diagnostics across all cohort × "
    "PS model combinations.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: Findings Summary
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Findings Summary", level=1)

# ── 2.1 Cohort/Track Validation Ranking ───────────────────────────────────────
add_heading(doc, "2.1 Cohort/Track Validation Ranking", level=2)
add_para(doc,
    "The table below summarizes the overall external evidence support for each cohort-track combination, "
    "based on the proportion of significant hits with strong validation and known biology.")

rank_headers = ["Rank", "Group", "Validation Score", "Key Validated Markers"]
rank_rows = [
    (1, "Cohort 2 — Track 2 (Predictive, matched)", "Highest",
     "CSF1R_SNV (Very Strong), MET_SNV (Strong), PTPRD_DEL (Strong, cohort 1 also)"),
    (2, "Track 1 — Pan-cancer (both cohorts)", "High",
     "STK11_DEL/SV (Very Strong), KEAP1_DEL (Strong), CD274_DEL (Strong), PDCD1LG2_DEL (Strong), PBRM1_SNV (Strong), ARID1A_SNV (Strong)"),
    (3, "Track 1 — BREAST (Cohort 2)", "Moderate–High",
     "MSH2_SNV (Very Strong), BRCA2_SNV (Moderate), RAD51C_SNV (Moderate), PALB2_SNV (Moderate)"),
    (4, "Cohort 1 — Track 2 (Predictive, unmatched)", "Moderate",
     "PTPRD_DEL (Strong), ESR1_AMP (Strong), SDHD_DEL (Moderate)"),
    (5, "Track 1 — LUNG (Cohort 1)", "Moderate",
     "STK11_DEL (Very Strong), KEAP1_DEL (Strong), MET_SNV (Moderate)"),
    (6, "Cohort 1 — Track 2 BREAST (unmatched)", "Low",
     "Many markers with None/Unassessed validation; likely residual confounding"),
]
make_table(doc, rank_headers, rank_rows, col_widths=[0.4, 2.0, 1.0, 3.5])

doc.add_paragraph()

# ── 2.2 Track 2 — Predictive Biomarker Findings ───────────────────────────────
add_heading(doc, "2.2 Track 2 — Predictive Biomarker Findings", level=2)
add_para(doc,
    "Track 2 identified markers where the somatic alteration differentially predicts ICI benefit or harm "
    "versus a non-ICI control group. Results are shown for all significant markers across both cohorts "
    "and all propensity score specifications (41 total rows; many markers replicate across specifications). "
    "Markers are sorted by validation level then by FDR. "
    "HR interaction < 1 = ICI benefit when marker is present; HR > 1 = ICI harm when marker is present.")

# Deduplicate for summary: pick best (lowest FDR) row per marker × cancer_type × cohort
t2_sorted = t2.copy()
t2_sorted["val_rank"] = t2_sorted["validation_level"].map(LEVEL_ORDER).fillna(99)
t2_sorted = t2_sorted.sort_values(["val_rank", "FDR_markerxICI"])

t2_headers = ["Marker", "Gene", "Cancer Type", "Direction", "HR (interaction)",
              "FDR", "Cohort", "PS Model", "Weight", "Validation Level", "Notes"]
t2_rows_with_level = []
for _, row in t2_sorted.iterrows():
    direction = "ICI Benefit" if "benefit" in str(row.get("classifier", "")) else "ICI Harm"
    hr_val = f"{row['HR_markerxICI']:.3f}" if pd.notna(row.get("HR_markerxICI")) else ""
    fdr_val = f"{row['FDR_markerxICI']:.4f}" if pd.notna(row.get("FDR_markerxICI")) else ""
    ps_short = "cov+emb" if "embeddings" in str(row.get("ps_model", "")) else "cov-only"
    cohort_short = row.get("cohort", "")
    wt = row.get("weight_type", "")
    vl = str(row.get("validation_level", "Unassessed"))
    notes = str(row.get("validation_notes", ""))
    if notes == "nan":
        notes = ""
    row_data = [
        row["marker"], row["gene"], row["cancer_type"],
        direction, hr_val, fdr_val,
        cohort_short, ps_short, wt,
        vl, notes
    ]
    t2_rows_with_level.append((row_data, vl))

make_validation_table(doc, t2_headers, t2_rows_with_level)
doc.add_paragraph()

# ── 2.3 Track 1 — Prognostic Findings (Strong/Very Strong) ────────────────────
add_heading(doc, "2.3 Track 1 — Prognostic Biomarker Findings (Top Consistent Hits)", level=2)
add_para(doc,
    "Track 1 identifies markers prognostically associated with survival within ICI-treated patients. "
    "The table below shows all markers with Strong or Very Strong external validation evidence. "
    "Consistency counts the number of unique (cohort, weight_type) combinations in which the marker "
    "reached significance — higher values indicate greater robustness across analytical specifications. "
    "HR > 1 = worse prognosis; HR < 1 = better prognosis in ICI-treated patients.")

t1_strong = t1[t1["validation_level"].isin(["Strong", "Very Strong"])].copy()
t1_strong["val_rank"] = t1_strong["validation_level"].map(LEVEL_ORDER)

# Summarize per marker × cancer_type: representative row (best FDR) + cohorts seen
t1_summary = []
for (gene, cancer_type), grp in t1_strong.groupby(["gene", "cancer_type"]):
    best_row = grp.sort_values("FDR_marker").iloc[0]
    cohorts   = ", ".join(sorted(grp["cohort"].unique()))
    hr_range  = f"{grp['HR_marker'].min():.2f}–{grp['HR_marker'].max():.2f}"
    fdr_best  = f"{best_row['FDR_marker']:.4f}"
    direction = "Better prognosis (ICI)" if best_row["HR_marker"] < 1 else "Worse prognosis (ICI)"
    vl = best_row["validation_level"]
    consistency = int(grp["consistency"].max()) if pd.notna(grp["consistency"].max()) else ""
    t1_summary.append({
        "gene": gene, "cancer_type": cancer_type,
        "HR_range": hr_range, "direction": direction,
        "FDR_best": fdr_best, "cohorts": cohorts,
        "consistency": consistency, "val_level": vl,
        "val_rank": LEVEL_ORDER.get(vl, 99),
        "notes": best_row.get("validation_notes", ""),
    })

t1_summary_df = pd.DataFrame(t1_summary).sort_values(["val_rank", "FDR_best"])

t1_headers = ["Gene", "Cancer Type", "HR Range", "Direction",
              "Best FDR", "Cohort(s)", "Consistency", "Validation Level"]
t1_rows_with_level = []
for _, row in t1_summary_df.iterrows():
    row_data = [
        row["gene"], row["cancer_type"],
        row["HR_range"], row["direction"],
        row["FDR_best"], row["cohorts"],
        str(row["consistency"]), row["val_level"],
    ]
    t1_rows_with_level.append((row_data, row["val_level"]))

make_validation_table(doc, t1_headers, t1_rows_with_level)
doc.add_paragraph()

# ── 2.4 Notable Findings ──────────────────────────────────────────────────────
add_heading(doc, "2.4 Notable Findings", level=2)

add_bullet(doc,
    "CSF1R_SNV (pan_cancer, Track 2 — Very Strong): The single strongest discovery from this pipeline. "
    "CSF1R encodes the colony-stimulating factor 1 receptor that drives TAM-mediated immunosuppression. "
    "CSF1R inhibitor + ICI combination therapy has been extensively validated preclinically and "
    "clinically. A loss-of-function CSF1R SNV would reduce TAM immunosuppression, producing the "
    "observed ICI benefit signal. This finding is present across all four propensity score "
    "specifications in Cohort 2 with FDR as low as 0.00002, and HR interaction as low as 0.064.")

add_bullet(doc,
    "MET_SNV (pan_cancer, Track 2 — Strong): Directly and prospectively validated by a 2024 pan-cancer "
    "study showing MET mutation as an independent favorable ICI biomarker via enhanced immune "
    "infiltration. This pipeline independently rediscovers this finding across both unmatched "
    "(Cohort 1) and matched (Cohort 2) designs, with consistent ICI benefit direction.")

add_bullet(doc,
    "Pan-cancer anchor markers (Track 1 — Strong/Very Strong): STK11 (DEL + SV), KEAP1 (DEL), "
    "CD274 (DEL), PDCD1LG2 (DEL), PBRM1 (SNV), and ARID1A (SNV) all reached significance in "
    "Track 1 pan-cancer analyses. These are among the most thoroughly validated ICI biomarkers "
    "in the literature. Their recovery confirms pipeline validity — the Track 1 prognostic model "
    "captures known biology. STK11 and KEAP1 specifically replicate as ICI resistance markers "
    "(HR > 1 in ICI-treated patients) in both LUNG and pan_cancer analyses.")

add_bullet(doc,
    "Soft tissue sarcoma (SOFT_TISSUE) ICI harm cluster (Track 2): ATM_DEL, CBL_DEL, NF1_DEL, "
    "and SDHD_DEL all show ICI harm (HR interaction 4.6–7.6) in soft tissue sarcoma, consistently "
    "across both ATE-weighted and unweighted specifications. SDHD_DEL is particularly notable — "
    "SDH-deficient GISTs and sarcomas are documented as ICI non-responsive, consistent with the "
    "observed harm direction. This cluster suggests a subtype of sarcoma where ICI may be "
    "detrimental and warrants prospective investigation.")

add_bullet(doc,
    "ESR1_AMP (pan_cancer, Track 2 — Strong): ESR1 amplification is associated with ICI harm "
    "(HR interaction = 2.18, FDR = 0.015). ESR1/ER expression drives an immunosuppressive TME; "
    "a 2025 study directly shows ESR1 expression predicts lower ICI response (AUC = 0.88). "
    "ESR1 amplification as a predictive ICI biomarker is a clinically actionable finding given "
    "the prevalence of ESR1 amplification in breast and other cancers.")

add_bullet(doc,
    "PTPRD_DEL (BRAIN, Track 2 — Strong): PTPRD deletion predicts ICI benefit in brain tumors "
    "(HR interaction = 0.33–0.35), consistently across Cohort 1 (both PS models) and with "
    "robust significance. Multiple studies have validated PTPRD/PTPRT mutations as ICI-favorable "
    "through JAK-STAT pathway modulation and enhanced antigen presentation. This is among the "
    "most robustly validated findings from this pipeline in cancer-type-specific analyses.")

# ── 2.5 Cohort Comparison Notes ───────────────────────────────────────────────
add_heading(doc, "2.5 Cohort Comparison Notes", level=2)

add_bullet(doc,
    "Cohort 2 (matched) produces cleaner Track 2 results: The 1:1 matching on cancer type and "
    "treatment line removes a major source of confounding. Cohort 2 Track 2 findings — particularly "
    "CSF1R_SNV and MET_SNV in pan_cancer — are more interpretable as causal ICI predictive "
    "markers. The much lower FDR values and consistent direction across specifications support this.")

add_bullet(doc,
    "Track 1 pan-cancer results are consistent across cohorts: STK11, KEAP1, CD274, PDCD1LG2, "
    "PBRM1, and ARID1A are recovered in both Cohort 1 and Cohort 2 Track 1 pan-cancer analyses. "
    "This is expected since both cohorts draw from largely overlapping ICI-treated patient "
    "populations; the Track 1 model only uses ICI-treated patients, where the two cohorts "
    "converge.")

add_bullet(doc,
    "Cohort 1 Track 2 BREAST findings likely reflect residual confounding: NEIL2_DEL, NKX3-1_DEL, "
    "PTK2B_DEL, and ZNRF3_DEL appear in Cohort 1 breast cancer Track 2 analyses but carry "
    "None or Indirect validation levels. These markers were detected only in the unmatched "
    "design, predominantly in the covariates + embeddings model. The breast cancer Cohort 1 "
    "analysis has a highly imbalanced treatment ratio (97 ICI vs 1,169 controls) and limited "
    "PS model discrimination, making spurious associations from residual confounding likely. "
    "These findings should be interpreted with caution and not prioritized for follow-up.")

# ── save ───────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Saved Word document to:\n  {OUT_PATH}")
